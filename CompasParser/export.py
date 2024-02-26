import pandas as pd
import xlsxwriter
from django.http import HttpResponse
import requests
from .models import NewsSource, ParsedNews
from datetime import datetime, date

from docx import Document

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import locale
from babel.dates import format_date
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Cm

from django.contrib.auth.decorators import permission_required


# функция экспорта в excel, не трогать - убьет!
def export_to_excel(request):
    if request.method == 'POST':
        selected_news_ids_str = request.POST.get('selected_news')
        
        # Разбиваем строку на список ID, разделенных запятыми
        selected_news_ids = [int(id_str) for id_str in selected_news_ids_str.split(',') if id_str]

        # Получите выбранные новости из базы данных (используйте вашу модель)
        selected_news = ParsedNews.objects.filter(id__in=selected_news_ids)

        # Update the 'selected' attribute to True
        selected_news.update(selected=True)
        
        selected_news_with_source = selected_news.select_related('source')
        # Создайте DataFrame с выбранными новостями
        data = {
            'Дата публикации': [news.date_published.strftime('%d.%m.%y') if news.date_published else '' for news in selected_news_with_source],
            'Название канала/страницы/паблика и ссылка на канал': [f"{news.source.name}\n {news.source.url}" if news.source else '' for news in selected_news_with_source],
            #'Название канала/страницы/паблика и ссылка на канал': [news.link for news in selected_news ],
            'Тема публикации/поста\n\n(например, телефонные мошенники)': ['' for news in selected_news_with_source],
            'Просмотры': [news.views if news.views is not None else 1 for news in selected_news_with_source],
            'Реакции(любые)': [(news.likes or 0) + (news.reposts or 0) + (news.comments or 0) for news in selected_news_with_source],
            #'Реакции(любые)': [news.likes + news.reposts + news.comments if all([news.likes, news.reposts, news.comments]) else 0 for news in selected_news_with_source],
            'Гиперссылка на материал\n(если есть)': [news.link for news in selected_news_with_source],
            'Примечание для больших групп: Общее количество подписчиков': [news.source.members_count if news.source.members_count else '' for news in selected_news_with_source],
            'Примечание 2': ['' for news in selected_news_with_source],
            'Примечание 3': ['' for news in selected_news_with_source],
        }

        df = pd.DataFrame(data)

        # Создайте Excel-файл с данными DataFrame
        response = HttpResponse(content_type='application/ms-excel')
        response['Content-Disposition'] = 'attachment; filename="exported_news.xlsx"'

        # Создайте экземпляр XlsxWriter с 'remove_timezone' опцией
        workbook = xlsxwriter.Workbook(response, {'in_memory': True, 'remove_timezone': True})

        # Создайте лист в файле Excel
        worksheet = workbook.add_worksheet('News')
      
        # Формат для ячеек с черными границами
        border_format = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True})
        
        # Формат для шапки таблицы с серым фоном
        header_format = workbook.add_format({'border': 1, 'bold': True, 'bg_color': '#D3D3D3', 'align': 'center', 'valign': 'vcenter', 'text_wrap': True})
        
        
        # Объединение ячеек и добавление черных границ для шапки
        worksheet.merge_range('A1:A2', 'Дата\nпубликации', header_format)
        worksheet.merge_range('B1:B2', 'Название канала/\nстраницы/паблика и\n ссылка на канал', header_format)
        worksheet.merge_range('C1:C2', 'Тема публикации/поста\n\n(например, телефонные\n мошенники)', header_format)
        # высота ячейки заданная
        worksheet.set_row(1, 60)
        worksheet.merge_range('D1:E1', 'ПАРАМЕТРЫ (за\n исключением групп ЦУР)', header_format)
        worksheet.set_row(0, 60)
        worksheet.write('D2', 'Просмотры', header_format)
        worksheet.write('E2', 'Реакции/\n(любые)', header_format)
        worksheet.merge_range('F1:F2', 'Гиперссылка на материал\n(если есть)', header_format)
        worksheet.merge_range('G1:G2', 'Примечание для\n больших групп:\nОбщее кол-во\nподписчиков', header_format)
        worksheet.merge_range('H1:H2', 'Примечание 2', header_format)
        worksheet.merge_range('I1:I2', 'Примечание 3', header_format)
        # Запишите данные DataFrame на лист
        for i, column_name in enumerate(df.columns):
            worksheet.write(1, i, column_name, border_format)
        
        # Определите ширину столбцов на основе максимальной длины содержимого в каждом столбце
        for i, column_name in enumerate(df.columns):
            max_length = max(df[column_name].astype(str).apply(len).max(), len(column_name))
            worksheet.set_column(1, i, max_length)  # Добавьте дополнительное пространство

        # закрашивает шапку серым
        for i, column_name in enumerate(df.columns):
            worksheet.write(1, i, column_name, header_format)
            
        # центрирует значения в таблице, вклоючает перенос строк  
        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                worksheet.write(i + 2, j, value, border_format)
                
                
        # Create an URL format for the hyperlink cells
        url_format = workbook.add_format({'color': 'blue', 'underline': 1, 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True})

        # Iterate through all cells in the DataFrame
        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                if "http" in str(value) or "www." in str(value):  # Check if the cell contains a URL
                    worksheet.write_url(i + 2, j, value, url_format, string=str(value))
                else:
                    worksheet.write(i + 2, j, value, border_format)

#         url_format = workbook.add_format({'color': 'blue', 'underline': 1})


#         # Write the "Название канала/страницы/паблика и ссылка на канал" column with URLs
#         for i, (channel_name, channel_url) in enumerate(zip(df['Название канала/страницы/паблика и ссылка на канал'], df['Гиперссылка на материал\n(если есть)'])):
#             if channel_name and channel_url:
#                 worksheet.write_url(i + 2, 1, channel_url, url_format, string=channel_name)

        # Set autofilter for all columns
        worksheet.autofilter(1, 0, len(df) + 1, len(df.columns) - 1)

        # Закройте книгу
        workbook.close()

        return response
    

# def export_to_word(request):
#     if request.method == 'POST':
#         selected_news_ids_str = request.POST.get('selected_news_word')
#         # Разбиваем строку на список ID, разделенных запятыми
#         selected_news_ids = [int(id_str) for id_str in selected_news_ids_str.split(',') if id_str]

#         # Получите выбранные новости из базы данных (используйте вашу модель)
#         selected_news = ParsedNews.objects.filter(id__in=selected_news_ids)

#         # Создаем документ Word
#         document = Document()
        
#         # Устанавливаем размеры полей (в сантиметрах)
#         left_margin = Cm(1.5)
#         right_margin = Cm(1.25)
#         top_margin = Cm(1.25)
#         bottom_margin = Cm(0.75)

#         sections = document.sections
#         for section in sections:
#             section.left_margin = left_margin
#             section.right_margin = right_margin
#             section.top_margin = top_margin
#             section.bottom_margin = bottom_margin
 
#         # # Вставка пустого абзаца с отступом слева
#         # empty_paragraph = document.add_paragraph()
#         # empty_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
#         # empty_paragraph.paragraph_format.left_indent = Pt(30)  # Отступ в тысячных долях дюйма (1,5 см)

#         # Вставка картинки
#         img_path = '/var/www/compas/static/info.jpg'  # Укажите путь к вашей картинке
#         document.add_picture(img_path)  # Изменьте ширину и высоту по своему усмотрению
#         document.add_paragraph()
        
#         # Создаем текущую дату
#         current_date = datetime.now()

#         # Форматируем дату на русском языке
#         formatted_date = format_date(current_date, format="d MMMM y 'г.'", locale='ru_RU')
#         # Добавляем отформатированную дату в документ
#         date_paragraph = document.add_paragraph(formatted_date)
#         date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         date_run = date_paragraph.runs[0]
#         date_run.font.size = Pt(14)
#         date_run.bold = True
#         date_run.font.name = 'Arial'

#         name = document.add_paragraph("Российская Федерация")  # Добавляем пустой абзац после источника и даты
#         name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         name_run = name.runs[0]
#         name_run.font.size = Pt(16)
#         name_run.font.bold = True
#         name_run.font.name = 'Arial'
#         name_run.font.color.rgb = RGBColor(18, 98, 73)  # Черный цвет
            
#         # Добавляем выбранные новости в документ
#         for news in selected_news:
#             # Добавляем заголовок новости и содержимое
#             title_paragraph = document.add_heading(clean_text(re.match(r'([^.!?]+[.!?])', news.content).group(0).strip()), level=1)
#             title_run = title_paragraph.runs[0]
#             title_run.font.size = Pt(12)
#             title_run.font.bold = True
#             title_run.font.name = 'Arial'
#             title_run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
#             title_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#             # Создаем абзац для источника и даты
#             source_and_date_paragraph = document.add_paragraph()
            
            
#             date_run = source_and_date_paragraph.add_run(news.date_published.strftime('%d.%m.%Y'))
#             date_run.font.size = Pt(12)
#             date_run.font.bold = True
#             date_run.font.name = 'Arial'
#             date_run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
#             source_and_date_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
#             source_and_date_paragraph.add_run(", ")
            
    
#             # Вставляем источник и дату в этот абзац
#             source_run = source_and_date_paragraph.add_run(news.source.name)
#             source_run.font.size = Pt(12)
#             source_run.font.bold = True
#             source_run.font.name = 'Arial'
#             source_run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
#             source_and_date_paragraph.add_run(" ")

            
            
            
#             source_run = source_and_date_paragraph.add_run(news.link)
#             source_run.font.size = Pt(12)
#             source_run.font.bold = True
#             source_run.font.name = 'Arial'
#             source_run.font.color.rgb = RGBColor(0, 0, 255)  # Черный цвет
#             source_and_date_paragraph.add_run(" ")

#             #document.add_paragraph()  # Добавляем пустой абзац после источника и даты

#             content_paragraph = document.add_paragraph()
#             content_paragraph.space_before = Pt(12.5)  # Устанавливаем отступ перед абзацем в 1,25 см
#             content_run = content_paragraph.add_run(clean_text(news.content))
#             content_run.font.size = Pt(12)
#             content_run.font.name = 'Arial'
#             content_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

#         response = HttpResponse(content_type='application/msword')
#         response['Content-Disposition'] = 'attachment; filename="selected_news.docx"'

#         # Сохраняем документ в HttpResponse
#         document.save(response)
#         return response
  
    

    
    
    
    
    
import qrcode
from django.http import HttpResponse
from docx import Document
from docx.shared import Cm, Pt, RGBColor

from .models import ParsedNews
from django.utils.timezone import localtime


import qrcode
from io import BytesIO
from django.http import HttpResponse
from docx import Document
from docx.shared import Cm
from .models import ParsedNews

def generate_qr_code(image_data):
    # Создание QR-кода из данных
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(image_data)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white")

    # Сохранение QR-кода в память
    qr_stream = BytesIO()
    qr_img.save(qr_stream, format='PNG')
    qr_stream.seek(0)
    return qr_stream


@permission_required('compasparser.can_export_to_word')
def export_to_word(request):
    if request.method == 'POST':
        selected_news_ids_str = request.POST.get('selected_news_word')
        selected_news_ids = [int(id_str) for id_str in selected_news_ids_str.split(',') if id_str]
        selected_news = ParsedNews.objects.filter(id__in=selected_news_ids)

        # Создаем документ Word
        document = Document()
        
        left_margin = Cm(1.5)
        right_margin = Cm(1.25)
        top_margin = Cm(1.25)
        bottom_margin = Cm(0.75)

        sections = document.sections
        for section in sections:
            section.left_margin = left_margin
            section.right_margin = right_margin
            section.top_margin = top_margin
            section.bottom_margin = bottom_margin
        table = document.add_table(rows=1, cols=6)  # Увеличиваем количество столбцов на 1
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№'
        hdr_cells[1].text = 'Ссылка'
        hdr_cells[2].text = 'Заголовок'  # Новый столбец
        hdr_cells[3].text = 'Текст'
        hdr_cells[4].text = 'QR-код'
        hdr_cells[5].text = 'Тема'

        for index, news in enumerate(selected_news, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = news.link
            row_cells[2].text = clean_text(clean_text(re.match(r'([^.!?]+[.!?])', news.content).group(0).strip()))  # Заголовок новости
            row_cells[3].text = clean_text(news.content)

            # Генерация QR-кода и добавление его в документ
            qr_stream = generate_qr_code(news.link)
            row_cells[4].paragraphs[0].add_run().add_picture(qr_stream, width=Cm(3))

        # Сохраняем документ в HttpResponse
        response = HttpResponse(content_type='application/msword')
        response['Content-Disposition'] = 'attachment; filename="selected_news.docx"'
        document.save(response)
        return response


    
    
import re

def clean_text(text):
    # Удаление стикеров
    text = re.sub(r'🟢|🔜|🤔|<.*?>', '', text)
    # Удаление хэштегов полностью (вместе с символом #)
    text = re.sub(r'#\w+', '', text)
    text = re.sub(r'@\w+', '', text)
    # Удаление ссылок в скобках
    text = re.sub(r'\((?=.*https?://)[^)]*\)', '', text)
    # Удаление подчеркивания
    text = re.sub(r'_', '', text)
    # Оставляем только буквы, цифры, пробелы, знаки препинания, дефис, тире и символы в круглых скобках
    text = re.sub(r'[^\w\s.,?!:;()—–-]', '', text)
    # Удаление двойных пробелов
    text = re.sub(r'\s+', ' ', text)
    return text.strip()